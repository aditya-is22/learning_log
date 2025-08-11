from django.db.models import Q
from django.shortcuts import render, get_object_or_404, redirect
from .models import Topic, Entry
from .forms import TopicForm
from .forms import EntryForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse, JsonResponse
from django.template.loader import render_to_string
from django.utils import timezone
import json
import csv
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches


# Create your views here.

def index(request):
    """The home page for Learning Log"""
    return render(request, 'learning_logs/index.html')

@login_required
def topics(request):
    """Show all topics."""
    topics = Topic.objects.filter(owner=request.user).order_by('-date_added') # Order by date_added in descending order
    context = {'topics': topics}
    return render(request, 'learning_logs/topics.html', context)

@login_required
def topic(request, topic_id):
    """Show a single topic and all its entries."""
    # Make sure the topic belongs to the current user.
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    entries = topic.entry_set.order_by('-date_added')
    context = {'topic': topic, 'entries': entries}
    return render(request, 'learning_logs/topic.html', context)

@login_required
def new_topic(request):
    """Add a new topic."""
    if request.method != 'POST':
        # No data submitted; create a blank form
        form = TopicForm()
    else:
        # POST data submitted; process data.
        form = TopicForm(data=request.POST)
        if form.is_valid():
            new_topic = form.save(commit=False)
            new_topic.owner = request.user
            new_topic.save()
            messages.success(request, f'Topic "{new_topic.text}" has been created successfully.')
            return redirect('learning_logs:topics')  # Redirect to topics page

    context = {'form': form}
    return render(request, 'learning_logs/new_topic.html', context)

@login_required
def new_entry(request, topic_id):
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)

    if request.method != 'POST':
        form = EntryForm()
    else:
        form = EntryForm(data=request.POST)
        if form.is_valid():
            new_entry = form.save(commit=False)
            new_entry.topic = topic
            new_entry.save()
            messages.success(request, 'New entry has been added successfully.')
            return redirect('learning_logs:topic', topic_id=topic_id)

    context = {'topic': topic, 'form': form}
    return render(request, 'learning_logs/new_entry.html', context)

@login_required
def edit_entry(request, entry_id):
    entry = get_object_or_404(Entry, id=entry_id, topic__owner=request.user)
    topic = entry.topic

    if request.method != 'POST':
        form = EntryForm(instance=entry)
    else:
        form = EntryForm(instance=entry, data=request.POST)
        if form.is_valid():
            form.save()
            messages.success(request, 'Entry has been updated successfully.')
            return redirect('learning_logs:topic', topic_id=topic.id)

    context = {'entry': entry, 'topic': topic, 'form': form}
    return render(request, 'learning_logs/edit_entry.html', context)


@login_required
def search_results(request):
    """Show search results for topics and entries."""
    query = request.GET.get('q')
    topics = Topic.objects.none()
    entries = Entry.objects.none()

    if query:
        # Find all topics owned by the user that match the query
        topics = Topic.objects.filter(
            owner=request.user,
            text__icontains=query
        ).distinct()

        # Find all entries where the entry or its topic matches the query
        entries = Entry.objects.filter(
            Q(topic__owner=request.user),
            Q(text__icontains=query) | Q(topic__text__icontains=query)
        ).select_related('topic').distinct()

    # The context is passed to the template without any text processing
    context = {'query': query, 'topics': topics, 'entries': entries}
    return render(request, 'learning_logs/search_results.html', context)


@login_required
def delete_topic(request, topic_id):
    """Delete a topic and all its entries."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    
    if request.method == 'POST':
        topic_name = topic.text
        topic.delete()
        messages.success(request, f'Topic "{topic_name}" and all its entries have been deleted successfully.')
        return redirect('learning_logs:topics')
    
    # For GET requests, redirect to topics page
    return redirect('learning_logs:topics')


@login_required  
def delete_entry(request, entry_id):
    """Delete an entry."""
    entry = get_object_or_404(Entry, id=entry_id, topic__owner=request.user)
    topic = entry.topic
    
    if request.method == 'POST':
        entry.delete()
        messages.success(request, 'Entry has been deleted successfully.')
        return redirect('learning_logs:topic', topic_id=topic.id)
    
    # For GET requests, redirect to topic page
    return redirect('learning_logs:topic', topic_id=topic.id)


@login_required
def export_all_entries(request, format):
    """Export all entries for the current user in the specified format."""
    topics = Topic.objects.filter(owner=request.user).order_by('-date_added')
    
    if format == 'json':
        return export_json(topics, request.user.username)
    elif format == 'csv':
        return export_csv(topics, request.user.username)
    elif format == 'pdf':
        return export_pdf(topics, request.user.username)
    elif format == 'docx':
        return export_docx(topics, request.user.username)
    else:
        messages.error(request, 'Invalid export format.')
        return redirect('learning_logs:topics')


@login_required
def export_topic_entries(request, topic_id, format):
    """Export entries for a specific topic in the specified format."""
    topic = get_object_or_404(Topic, id=topic_id, owner=request.user)
    topics = [topic]  # Wrap in list for consistency with export functions
    
    if format == 'json':
        return export_json(topics, f"{request.user.username}_{topic.text}")
    elif format == 'csv':
        return export_csv(topics, f"{request.user.username}_{topic.text}")
    elif format == 'pdf':
        return export_pdf(topics, f"{request.user.username}_{topic.text}")
    elif format == 'docx':
        return export_docx(topics, f"{request.user.username}_{topic.text}")
    else:
        messages.error(request, 'Invalid export format.')
        return redirect('learning_logs:topic', topic_id=topic_id)


def export_json(topics, filename_prefix):
    """Export topics and entries to JSON format."""
    data = {
        'export_date': timezone.now().isoformat(),
        'topics': []
    }
    
    for topic in topics:
        entries = topic.entry_set.order_by('-date_added')
        topic_data = {
            'topic_name': topic.text,
            'date_added': topic.date_added.isoformat(),
            'entries': [
                {
                    'text': entry.text,
                    'date_added': entry.date_added.isoformat(),
                    'date_modified': entry.date_modified.isoformat()
                }
                for entry in entries
            ]
        }
        data['topics'].append(topic_data)
    
    response = HttpResponse(
        json.dumps(data, indent=2, ensure_ascii=False),
        content_type='application/json'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_learning_log.json"'
    return response


def export_csv(topics, filename_prefix):
    """Export topics and entries to CSV format."""
    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_learning_log.csv"'
    
    writer = csv.writer(response)
    writer.writerow(['Topic', 'Entry Text', 'Date Added', 'Date Modified'])
    
    for topic in topics:
        entries = topic.entry_set.order_by('-date_added')
        for entry in entries:
            writer.writerow([
                topic.text,
                entry.text,
                entry.date_added.strftime('%Y-%m-%d %H:%M:%S'),
                entry.date_modified.strftime('%Y-%m-%d %H:%M:%S')
            ])
    
    return response


def export_pdf(topics, filename_prefix):
    """Export topics and entries to PDF format."""
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Title
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        alignment=1  # Center alignment
    )
    story.append(Paragraph("Learning Log Export", title_style))
    story.append(Spacer(1, 0.5*inch))
    
    # Topics and entries
    for topic in topics:
        # Topic title
        topic_style = ParagraphStyle(
            'TopicTitle',
            parent=styles['Heading2'],
            fontSize=16,
            spaceAfter=12,
            textColor='blue'
        )
        story.append(Paragraph(f"Topic: {topic.text}", topic_style))
        story.append(Paragraph(f"Created: {topic.date_added.strftime('%B %d, %Y')}", styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Entries
        entries = topic.entry_set.order_by('-date_added')
        if entries:
            for entry in entries:
                # Entry date
                entry_date_style = ParagraphStyle(
                    'EntryDate',
                    parent=styles['Normal'],
                    fontSize=10,
                    textColor='gray',
                    spaceAfter=6
                )
                story.append(Paragraph(f"Entry from {entry.date_added.strftime('%B %d, %Y at %I:%M %p')}", entry_date_style))
                
                # Entry text
                story.append(Paragraph(entry.text, styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
        else:
            story.append(Paragraph("No entries for this topic yet.", styles['Italic']))
        
        story.append(Spacer(1, 0.3*inch))
    
    doc.build(story)
    
    response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_learning_log.pdf"'
    return response


def export_docx(topics, filename_prefix):
    """Export topics and entries to Word document format."""
    doc = Document()
    
    # Title
    title = doc.add_heading('Learning Log Export', 0)
    title.alignment = 1  # Center alignment
    
    # Topics and entries
    for topic in topics:
        # Topic heading
        topic_heading = doc.add_heading(f"Topic: {topic.text}", level=1)
        doc.add_paragraph(f"Created: {topic.date_added.strftime('%B %d, %Y')}")
        
        # Entries
        entries = topic.entry_set.order_by('-date_added')
        if entries:
            for entry in entries:
                # Entry subheading
                entry_heading = doc.add_heading(f"Entry from {entry.date_added.strftime('%B %d, %Y at %I:%M %p')}", level=2)
                
                # Entry text
                doc.add_paragraph(entry.text)
                
                # Add space between entries
                doc.add_paragraph()
        else:
            doc.add_paragraph("No entries for this topic yet.", style='Intense Quote')
        
        # Add page break between topics (except for the last one)
        if topic != topics[-1]:
            doc.add_page_break()
    
    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="{filename_prefix}_learning_log.docx"'
    return response